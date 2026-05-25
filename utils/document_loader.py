"""
文档加载器模块 - 支持多种格式的文档加载和OCR识别

支持的格式:
- PDF: 文本层提取 + 图片OCR
- Word (.docx): 段落文字 + 图片OCR
- 图片 (jpg/png): OCR识别
- TXT: 纯文本加载
"""
import os
from typing import Optional
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from rag.multimodal_ocr import ocr_service
from utils.logger_handler import logger


def load_pdf(filepath: str, password: Optional[str] = None) -> list[Document]:
    """
    加载PDF文档：同时提取文字层和图片OCR内容
    
    Args:
        filepath: PDF文件路径
        password: PDF密码（可选）
        
    Returns:
        Document列表
    """
    documents = []
    
    # 1. 提取文本层
    try:
        text_docs = PyPDFLoader(filepath, password).load()
        if text_docs:
            logger.info(f"[PDF] {filepath} 提取文本层: {len(text_docs)}页")
            documents.extend(text_docs)
    except Exception as e:
        logger.warning(f"[PDF] {filepath} 文本提取失败: {e}")
    
    # 2. 提取并OCR图片
    try:
        image_docs = _extract_images_from_pdf(filepath)
        if image_docs:
            logger.info(f"[PDF] {filepath} OCR图片: {len(image_docs)}个内容块")
            documents.extend(image_docs)
    except Exception as e:
        logger.warning(f"[PDF] {filepath} 图片提取失败: {e}")
    
    if documents:
        logger.info(f"[PDF] {filepath} 共生成 {len(documents)} 个文档块")
    else:
        logger.warning(f"[PDF] {filepath} 未提取到内容")
    
    return documents


def _extract_images_from_pdf(pdf_path: str) -> list[Document]:
    """从PDF中提取图片并OCR识别（复用 multimodal_ocr 服务）"""
    documents = []

    try:
        # 使用 multimodal_ocr 服务的 PDF 页面识别功能
        results = ocr_service.recognize_pdf_pages(pdf_path)

        for page_num, text in results:
            if text and len(text.strip()) > 10:  # 过滤短文本
                doc = Document(
                    page_content=f"【第{page_num}页扫描内容】\n{text}",
                    metadata={
                        "source": pdf_path,
                        "page": page_num,
                        "type": "pdf_scan_ocr"
                    }
                )
                documents.append(doc)

        if documents:
            logger.info(f"[PDF] {pdf_path} OCR扫描 {len(documents)} 个页面")
        return documents

    except Exception as e:
        logger.error(f"[PDF] {pdf_path} OCR失败: {e}", exc_info=True)
        return []


def load_docx(filepath: str) -> list[Document]:
    """
    加载Word文档：提取文字和图片OCR内容
    
    Args:
        filepath: Word文件路径(.docx)
        
    Returns:
        Document列表
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.error("[Word] 请安装: pip install python-docx")
        return []
    
    documents = []
    
    try:
        doc = DocxDocument(filepath)
        logger.info(f"[Word] {filepath} 开始解析，{len(doc.paragraphs)}个段落")
        
        # 1. 提取文字
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        if text_parts:
            full_text = "\n\n".join(text_parts)
            documents.append(Document(
                page_content=full_text,
                metadata={"source": filepath, "type": "docx_text"}
            ))
            logger.info(f"[Word] {filepath} 文字: {len(full_text)}字符")
        
        # 2. 提取并OCR图片
        image_docs = _extract_images_from_docx(filepath)
        if image_docs:
            logger.info(f"[Word] {filepath} OCR图片: {len(image_docs)}个")
            documents.extend(image_docs)
        
        if documents:
            logger.info(f"[Word] {filepath} 共{len(documents)}个文档块")
        else:
            logger.warning(f"[Word] {filepath} 未提取到内容")
        
        return documents
        
    except Exception as e:
        logger.error(f"[Word] {filepath} 失败: {e}", exc_info=True)
        return []


def _extract_images_from_docx(docx_path: str) -> list[Document]:
    """从Word文档中提取图片并OCR识别"""
    try:
        import zipfile
        import tempfile
    except ImportError:
        logger.error("[Word图片] 请安装: pip install python-docx")
        return []

    documents = []

    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            # 查找word/media/下的图片
            image_files = [
                f for f in zip_ref.namelist()
                if f.startswith('word/media/') and
                f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))
            ]

            if not image_files:
                logger.debug(f"[Word图片] {docx_path} 无图片")
                return []

            logger.debug(f"[Word图片] {docx_path} 找到{len(image_files)}张图片")

            for idx, img_path_in_zip in enumerate(image_files, start=1):
                try:
                    # 读取图片数据
                    img_data = zip_ref.read(img_path_in_zip)

                    # 确定扩展名
                    ext = os.path.splitext(img_path_in_zip)[1].lower()
                    suffix = '.jpg' if ext not in ['.png'] else ext

                    # 保存为临时文件
                    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                        tmp.write(img_data)
                        tmp_path = tmp.name

                    try:
                        # 使用 multimodal_ocr 服务识别图片
                        text = ocr_service.recognize_image(tmp_path)
                        if text and len(text.strip()) > 10:
                            documents.append(Document(
                                page_content=f"【Word文档内图片{idx}】\n{text}",
                                metadata={
                                    "source": docx_path,
                                    "image_name": os.path.basename(img_path_in_zip),
                                    "type": "docx_image_ocr"
                                }
                            ))
                    finally:
                        if os.path.exists(tmp_path):
                            os.remove(tmp_path)

                except Exception as e:
                    logger.debug(f"[Word图片] 图片{idx} OCR失败: {e}")
                    continue

        return documents

    except Exception as e:
        logger.error(f"[Word图片] {docx_path} 失败: {e}", exc_info=True)
        return []


def load_image(filepath: str) -> list[Document]:
    """
    加载图片文件：OCR识别文字
    
    Args:
        filepath: 图片路径(jpg/jpeg/png)
        
    Returns:
        Document列表
    """
    try:
        logger.info(f"[图片OCR] {filepath}")
        text = ocr_service.recognize_image(filepath)
        
        if not text:
            logger.warning(f"[图片OCR] {filepath} 未识别到文字")
            return []
        
        doc = Document(
            page_content=text,
            metadata={"source": filepath, "type": "image_ocr"}
        )
        
        logger.info(f"[图片OCR] {filepath} 成功: {len(text)}字符")
        return [doc]
        
    except Exception as e:
        logger.error(f"[图片OCR] {filepath} 失败: {e}", exc_info=True)
        return []


def load_txt(filepath: str) -> list[Document]:
    """
    加载TXT文本文件
    
    Args:
        filepath: TXT文件路径
        
    Returns:
        Document列表
    """
    try:
        docs = TextLoader(filepath, encoding="utf-8").load()
        logger.info(f"[TXT] {filepath} 加载成功")
        return docs
    except Exception as e:
        logger.error(f"[TXT] {filepath} 失败: {e}")
        return []
